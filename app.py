# app.py  -- Stage 4: RAG using local reference files (simple TF-IDF retrieval)
import streamlit as st
from docx import Document
import io
from PyPDF2 import PdfReader
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import os

# -----------------------------
# Utilities: read reference files (PDF / DOCX)
# -----------------------------
def read_pdf_bytes(file_bytes):
    reader = PdfReader(io.BytesIO(file_bytes))
    texts = []
    for page in reader.pages:
        try:
            texts.append(page.extract_text() or "")
        except Exception:
            texts.append("")
    return "\n".join(texts)

def read_docx_bytes_to_text(file_bytes):
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip() != ""]
    return "\n".join(paragraphs)

# -----------------------------
# Chunking function (split long text into smaller searchable chunks)
# -----------------------------
def chunk_text(text, chunk_size_paragraphs=3):
    # naive approach: split by newline paragraphs and group
    paras = [p.strip() for p in text.split("\n") if p.strip()]
    chunks = []
    for i in range(0, len(paras), chunk_size_paragraphs):
        chunk = " ".join(paras[i:i+chunk_size_paragraphs])
        chunks.append(chunk)
    # if no paragraphs, keep whole text as one chunk
    if not chunks and text.strip():
        chunks = [text.strip()]
    return chunks

# -----------------------------
# Build local retrieval index given uploaded reference files
# -----------------------------
def build_reference_index(ref_files):
    """
    ref_files: list of tuples (filename, bytes)
    returns:
      - corpus_chunks: list of text chunks
      - chunk_sources: list of dicts {filename, chunk_text, chunk_id}
      - vectorizer, matrix (tfidf)
    """
    corpus_chunks = []
    chunk_sources = []
    for fname, file_bytes in ref_files:
        f_lower = fname.lower()
        try:
            if fname.endswith(".pdf") or f_lower.endswith(".pdf"):
                text = read_pdf_bytes(file_bytes)
            else:
                # assume docx
                text = read_docx_bytes_to_text(file_bytes)
        except Exception as e:
            text = ""
        chunks = chunk_text(text, chunk_size_paragraphs=3)
        for idx, c in enumerate(chunks):
            chunk_id = f"{fname}::chunk_{idx+1}"
            corpus_chunks.append(c)
            chunk_sources.append({"filename": fname, "chunk_id": chunk_id, "text": c})
    # build TF-IDF matrix
    if corpus_chunks:
        vect = TfidfVectorizer(stop_words="english", max_df=0.9)
        tfidf_matrix = vect.fit_transform(corpus_chunks)
    else:
        vect = None
        tfidf_matrix = None
    return corpus_chunks, chunk_sources, vect, tfidf_matrix

# -----------------------------
# Retrieve top matching reference chunk(s) for a query
# -----------------------------
def retrieve_law_passages(query, vectorizer, tfidf_matrix, chunk_sources, top_k=1):
    """
    returns list of best matches: [{'filename', 'chunk_id', 'text', 'score'}]
    """
    if vectorizer is None or tfidf_matrix is None or not chunk_sources:
        return []
    qv = vectorizer.transform([query])
    scores = cosine_similarity(qv, tfidf_matrix)[0]  # array of similarity scores
    top_indices = scores.argsort()[::-1][:top_k]
    results = []
    for i in top_indices:
        results.append({
            "filename": chunk_sources[i]["filename"],
            "chunk_id": chunk_sources[i]["chunk_id"],
            "text": chunk_sources[i]["text"],
            "score": float(scores[i])
        })
    return results

# -----------------------------
# Read uploaded user docx (same as before)
# -----------------------------
def read_docx_bytes(file_bytes):
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip() != ""]
    return "\n".join(paragraphs)

# -----------------------------
# Simple red-flag checks (same as Stage 2/3)
# -----------------------------
def check_red_flags(text):
    issues = []
    lower_text = text.lower()
    if "uae federal court" in lower_text:
        issues.append({
            "issue": "Jurisdiction clause mentions 'UAE Federal Court' instead of 'ADGM Courts'",
            "severity": "High",
            "suggestion": "Change to 'ADGM Courts' as per ADGM Companies Regulations."
        })
    if " may " in lower_text:
        issues.append({
            "issue": "Document uses 'may' which can be ambiguous",
            "severity": "Medium",
            "suggestion": "Consider using 'shall' for stronger legal binding."
        })
    if "signature" not in lower_text and "signed by" not in lower_text:
        issues.append({
            "issue": "No signature section found",
            "severity": "High",
            "suggestion": "Add a signature section to validate the document."
        })
    return issues

# -----------------------------
# Add comments + retrieved law passage to a reviewed docx and return bytes IO
# -----------------------------
def create_reviewed_docx_with_law(original_bytes, issues_with_retrieval):
    doc = Document(io.BytesIO(original_bytes))
    # Append review comments and retrieved law text
    if issues_with_retrieval:
        doc.add_paragraph("\n--- REVIEW COMMENTS ---")
        for item in issues_with_retrieval:
            # item: {'issue','severity','suggestion','retrievals': [{'filename','text','score'}, ...]}
            doc.add_paragraph(f"Issue: {item['issue']}")
            doc.add_paragraph(f"Severity: {item['severity']}")
            doc.add_paragraph(f"Suggestion: {item['suggestion']}")
            # add retrieved law passages if any
            if item.get("retrievals"):
                doc.add_paragraph("Relevant ADGM reference(s):")
                for r in item["retrievals"]:
                    # include filename and snippet (short)
                    snippet = r["text"][:600] + ("..." if len(r["text"]) > 600 else "")
                    doc.add_paragraph(f"- Source: {r['filename']} (score: {r['score']:.3f})")
                    doc.add_paragraph(snippet)
            doc.add_paragraph("")  # blank line
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out

# -----------------------------
# Simple doc type detection (same as before)
# -----------------------------
def detect_document_type(text):
    t = text.lower()
    if "articles of association" in t:
        return "Articles of Association"
    if "memorandum of association" in t:
        return "Memorandum of Association"
    if "register of members" in t:
        return "Register of Members and Directors"
    if "ubo" in t or "ultimate beneficial owner" in t:
        return "UBO Declaration Form"
    if "board resolution" in t:
        return "Board Resolution"
    return "Unknown / Other"

# -----------------------------
# Checklist
# -----------------------------
INCORP_CHECKLIST = [
    "Articles of Association",
    "Memorandum of Association",
    "Board Resolution",
    "UBO Declaration Form",
    "Register of Members and Directors"
]

# -----------------------------
# Streamlit UI
# -----------------------------
st.title("ADGM Corporate Agent — Stage 4 (RAG prototype)")
st.write("Upload ADGM reference files (PDF/DOCX) first, then upload the user documents to review.")

# Reference files uploader (the ADGM files you downloaded)
ref_uploads = st.file_uploader("Upload ADGM reference files (PDF or DOCX) — e.g. checklists, templates", accept_multiple_files=True, type=["pdf","docx"])

# Build index button
if ref_uploads:
    ref_files = []
    for f in ref_uploads:
        f.seek(0)
        ref_files.append((f.name, f.read()))
    st.success(f"{len(ref_files)} reference file(s) uploaded. Building index...")
    corpus_chunks, chunk_sources, vectorizer, tfidf_matrix = build_reference_index(ref_files)
    st.info(f"Reference index built with {len(corpus_chunks)} chunks from {len(ref_files)} files.")
else:
    corpus_chunks, chunk_sources, vectorizer, tfidf_matrix = [], [], None, None

st.markdown("---")
# User document uploader
uploaded_files = st.file_uploader("Upload .docx files to review (user documents)", accept_multiple_files=True, type=["docx"])

if uploaded_files:
    detected_types = []
    all_issues = []

    for uploaded in uploaded_files:
        uploaded.seek(0)
        raw = uploaded.read()
        text = read_docx_bytes(raw)
        doc_type = detect_document_type(text)
        detected_types.append(doc_type)

        st.write(f"**{uploaded.name}** — Detected as: {doc_type}")
        st.text(text[:400] + ("..." if len(text) > 400 else ""))

        # red flags
        issues = check_red_flags(text)
        issues_with_retrieval = []
        if issues:
            st.error("Red Flags Found:")
            for issue in issues:
                st.write(f"- **Issue:** {issue['issue']}")
                st.write(f"  **Severity:** {issue['severity']}")
                st.write(f"  **Suggestion:** {issue['suggestion']}")

                # retrieve relevant law passages for this issue (if index present)
                retrievals = retrieve_law_passages(issue['issue'] + " " + issue['suggestion'],
                                                   vectorizer, tfidf_matrix, chunk_sources, top_k=2)
                if retrievals:
                    st.write("  **Relevant ADGM reference (top):**")
                    for r in retrievals:
                        st.write(f"    - {r['filename']} (score {r['score']:.3f})")
                        short = r['text'][:400] + ("..." if len(r['text']) > 400 else "")
                        st.text(short)
                else:
                    st.write("  (No ADGM reference matched — upload ADGM reference files first.)")

                issues_with_retrieval.append({
                    "issue": issue['issue'],
                    "severity": issue['severity'],
                    "suggestion": issue['suggestion'],
                    "retrievals": retrievals
                })
        else:
            st.success("No obvious red flags found.")
        # create reviewed docx containing the retrieved law passages
        reviewed = create_reviewed_docx_with_law(raw, issues_with_retrieval)
        st.download_button(
            label=f"Download Reviewed {uploaded.name}",
            data=reviewed,
            file_name=f"Reviewed_{uploaded.name}",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        all_issues.append({"document": uploaded.name, "issues": issues_with_retrieval})

    # Missing docs check
    missing = [doc for doc in INCORP_CHECKLIST if doc not in detected_types]
    if missing:
        st.warning("Missing required documents:")
        for m in missing:
            st.write("- " + m)
    else:
        st.success("All required incorporation documents appear to be uploaded.")

    # JSON summary
    report = {
        "process": "Company Incorporation",
        "documents_uploaded": len(detected_types),
        "required_documents": len(INCORP_CHECKLIST),
        "missing_documents": missing,
        "detected_types": detected_types,
        "issues_found": all_issues
    }
    st.subheader("Summary")
    st.json(report)
else:
    st.info("Upload user .docx files to review (after uploading reference files).")
